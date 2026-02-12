import streamlit as st
import time
import logging
import html as html_module
from datetime import datetime
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from bs4 import BeautifulSoup
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import re
import json
from pathlib import Path
from urllib.parse import urlparse, urljoin
import xml.etree.ElementTree as ET
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import FormulaRule
from openpyxl.utils import get_column_letter

from detector import detect_problems, Problem, normalize_url
from validator import validate_results


# ===========================
# LOGGING
# ===========================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("quickwins")


# ===========================
# CONSTANTS
# ===========================
CRAWL_TIMEOUT = 12
MAX_PAGES = 40
MAX_INTERNAL_LINKS_PER_PAGE = 10
MAX_BROKEN_LINK_CHECKS = 180
MAX_BROKEN_LINKS_REPORTED = 50
MAX_SITEMAP_URLS = 6000
THIN_CONTENT_THRESHOLD = 250
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
LLM_MAX_TOKENS = 1800
LLM_TEMPERATURE = 0.2
HTTP_MAX_RETRIES = 3
HTTP_BACKOFF_FACTOR = 1  # seconds: 1s, 2s, 4s


# ===========================
# PAGE CONFIGURATION
# ===========================
st.set_page_config(
    page_title="Quick Wins - SEO Audit Tool",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ===========================
# API CONFIGURATION
# ===========================
try:
    GEMINI_API_KEY = st.secrets.get("GOOGLE_API_KEY", "")
    if GEMINI_API_KEY:
        genai.configure(api_key=GEMINI_API_KEY)
        GEMINI_AVAILABLE = True
    else:
        GEMINI_AVAILABLE = False
        logger.warning("GOOGLE_API_KEY not found in secrets")
except Exception as e:
    GEMINI_AVAILABLE = False
    logger.error("Failed to configure API: %s", e)


# ===========================
# CUSTOM CSS
# ===========================
st.markdown("""
<style>
    .stApp {
        background: #0f172a;
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f172a 0%, #020617 100%);
    }

    [data-testid="stMetricValue"] {
        font-size: 24px;
        color: #60a5fa;
        font-weight: 600;
    }

    [data-testid="stMetricLabel"] {
        color: #94a3b8;
        font-size: 13px;
        font-weight: 500;
    }

    .stButton>button {
        width: 100%;
        background: linear-gradient(90deg, #60a5fa 0%, #3b82f6 100%);
        color: white;
        font-weight: 600;
        border: none;
        padding: 10px 20px;
        border-radius: 6px;
        font-size: 15px;
        transition: all 0.3s ease;
    }

    .stButton>button:hover {
        transform: translateY(-1px);
        box-shadow: 0 6px 12px rgba(96, 165, 250, 0.3);
    }

    .stTextInput>div>div>input {
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        border: 1px solid rgba(96, 165, 250, 0.3);
        border-radius: 6px;
        padding: 8px;
        font-size: 14px;
    }

    h1 { color: #60a5fa; font-weight: 700; }
    h2, h3 { color: #e2e8f0; }

    .app-header {
        text-align: center;
        padding: 20px 0 30px 0;
        margin-bottom: 30px;
        border-bottom: 2px solid rgba(96, 165, 250, 0.2);
    }
    .app-title {
        font-size: 48px;
        font-weight: 700;
        color: #60a5fa;
        margin: 10px 0 5px 0;
        letter-spacing: -1px;
    }
    .app-subtitle {
        font-size: 17px;
        color: #94a3b8;
        font-weight: 400;
        line-height: 1.6;
        max-width: 600px;
        margin: 0 auto;
    }

    /* Audit header banner */
    .audit-banner {
        background: linear-gradient(135deg, #1e3a5f 0%, #1e293b 100%);
        padding: 28px 32px;
        border-radius: 12px;
        margin-bottom: 28px;
        border: 1px solid rgba(96, 165, 250, 0.2);
    }
    .audit-banner h2 {
        color: #f8fafc;
        font-size: 28px;
        font-weight: 700;
        margin: 0;
    }

    /* Info boxes (Executive Summary, Scope) */
    .info-box {
        background: rgba(30, 41, 59, 0.8);
        border: 1px solid rgba(96, 165, 250, 0.2);
        border-radius: 10px;
        padding: 24px 28px;
        margin-bottom: 24px;
        line-height: 1.8;
        color: #cbd5e1;
    }
    .info-box h3 {
        color: #93c5fd;
        font-size: 18px;
        font-weight: 600;
        margin: 0 0 12px 0;
    }
    .info-box p {
        margin: 0 0 8px 0;
        color: #cbd5e1;
        font-size: 15px;
    }

    /* Section headers */
    .section-header {
        font-size: 22px;
        font-weight: 700;
        margin: 32px 0 16px 0;
        color: #f8fafc;
    }
    .section-subtitle {
        color: #94a3b8;
        font-size: 14px;
        margin: -10px 0 20px 0;
    }

    /* Quick Wins container */
    .qw-container {
        background: linear-gradient(135deg, #1a2e1a 0%, #0f2818 50%, #162316 100%);
        border: 1px solid rgba(34, 197, 94, 0.3);
        border-radius: 14px;
        padding: 28px 28px 18px 28px;
        margin-bottom: 28px;
    }
    .qw-container-title {
        font-size: 22px;
        font-weight: 700;
        color: #4ade80;
        margin: 0 0 4px 0;
    }
    .qw-container-subtitle {
        color: #86efac;
        font-size: 14px;
        margin: 0 0 20px 0;
        opacity: 0.8;
    }

    /* Quick Win cards */
    .qw-card {
        background: rgba(34, 197, 94, 0.08);
        border-left: 4px solid #22c55e;
        border-radius: 8px;
        padding: 18px 22px;
        margin-bottom: 12px;
        transition: background 0.2s;
    }
    .qw-card:hover {
        background: rgba(34, 197, 94, 0.15);
    }
    .qw-number {
        display: inline-block;
        background: #22c55e;
        color: #0f172a;
        font-weight: 700;
        font-size: 13px;
        width: 24px;
        height: 24px;
        line-height: 24px;
        text-align: center;
        border-radius: 4px;
        margin-right: 10px;
    }
    .qw-title {
        font-size: 16px;
        font-weight: 600;
        color: #4ade80;
        display: inline;
    }
    .qw-desc {
        color: #cbd5e1;
        font-size: 14px;
        margin-top: 8px;
        padding-left: 34px;
        line-height: 1.5;
    }
    .qw-urls {
        color: #60a5fa;
        font-size: 12px;
        font-family: monospace;
        margin-top: 6px;
        padding-left: 34px;
        word-break: break-all;
    }
    .qw-more {
        color: #86efac;
        font-size: 12px;
        font-style: italic;
        padding-left: 34px;
        margin-top: 4px;
    }

    /* Error / Warning cards inside expanders */
    .issue-card {
        background: rgba(15, 23, 42, 0.6);
        border-radius: 8px;
        padding: 18px 22px;
        margin-bottom: 14px;
        border: 1px solid rgba(100, 116, 139, 0.2);
    }
    .issue-card h4 {
        color: #f8fafc;
        font-size: 15px;
        font-weight: 600;
        margin: 0 0 8px 0;
    }
    .issue-card p {
        color: #94a3b8;
        font-size: 14px;
        margin: 4px 0;
        line-height: 1.6;
    }
    .issue-urls {
        color: #60a5fa;
        font-size: 13px;
        font-family: monospace;
        margin-top: 8px;
        word-break: break-all;
    }
    .issue-urls a {
        color: #60a5fa;
        text-decoration: none;
    }
    .issue-urls a:hover {
        text-decoration: underline;
    }
    .issue-detail-label {
        color: #93c5fd;
        font-weight: 600;
        font-size: 13px;
        margin-top: 10px;
        margin-bottom: 2px;
    }
    .issue-detail-text {
        color: #cbd5e1;
        font-size: 14px;
        line-height: 1.6;
    }
    .more-in-excel {
        color: #f59e0b;
        font-size: 13px;
        font-style: italic;
        margin-top: 6px;
    }

    /* Next checks cards */
    .next-card {
        background: rgba(30, 41, 59, 0.5);
        border: 1px solid rgba(96, 165, 250, 0.15);
        border-radius: 8px;
        padding: 16px 20px;
        margin-bottom: 10px;
    }
    .next-card h4 {
        color: #93c5fd;
        font-size: 15px;
        font-weight: 600;
        margin: 0 0 6px 0;
    }
    .next-card p {
        color: #94a3b8;
        font-size: 14px;
        margin: 0;
        line-height: 1.5;
    }

    /* Action buttons row */
    .btn-excel > button {
        background: linear-gradient(90deg, #f59e0b 0%, #d97706 100%) !important;
        color: #0f172a !important;
        font-weight: 700 !important;
    }
    .btn-excel > button:hover {
        box-shadow: 0 6px 12px rgba(245, 158, 11, 0.3) !important;
    }
    .btn-new-audit > button {
        background: transparent !important;
        border: 2px solid #60a5fa !important;
        color: #60a5fa !important;
        font-weight: 600 !important;
    }
    .btn-new-audit > button:hover {
        background: rgba(96, 165, 250, 0.1) !important;
    }

    /* Expander styling */
    .streamlit-expanderHeader {
        background: rgba(30, 41, 59, 0.8) !important;
        border-radius: 8px !important;
        color: #f8fafc !important;
        font-weight: 600 !important;
    }

    /* Status badges */
    .status-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 12px;
        font-weight: 600;
        margin: 4px;
    }
    .status-connected {
        background-color: rgba(34, 197, 94, 0.2);
        color: #22c55e;
        border: 1px solid #22c55e;
    }
    .status-disconnected {
        background-color: rgba(239, 68, 68, 0.2);
        color: #ef4444;
        border: 1px solid #ef4444;
    }

    /* Hide Streamlit default footer */
    footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)


# ===========================
# PATHS
# ===========================
BASE_DIR = Path(__file__).parent
PROMPTS_DIR = BASE_DIR / "prompts"
PROMPT_BASIC = PROMPTS_DIR / "basic.md"


# ===========================
# HTTP SESSION WITH RETRY
# ===========================
def _create_http_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(
        total=HTTP_MAX_RETRIES,
        backoff_factor=HTTP_BACKOFF_FACTOR,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["HEAD", "GET"],
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    session.headers.update({"User-Agent": USER_AGENT})
    return session


http_session = _create_http_session()


# ===========================
# URL VALIDATION
# ===========================
_PRIVATE_HOSTNAMES = {"localhost", "127.0.0.1", "0.0.0.0", "::1"}
_PRIVATE_PREFIXES = ("192.168.", "10.", "172.16.", "172.17.", "172.18.",
                     "172.19.", "172.20.", "172.21.", "172.22.", "172.23.",
                     "172.24.", "172.25.", "172.26.", "172.27.", "172.28.",
                     "172.29.", "172.30.", "172.31.", "169.254.")


def validate_url(raw: str) -> tuple[bool, str]:
    """Validate and sanitize a URL. Returns (is_valid, clean_url_or_error)."""
    url = (raw or "").strip()
    if not url:
        return False, "URL is required."
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    try:
        parsed = urlparse(url)
    except Exception:
        return False, "Malformed URL."
    if parsed.scheme not in ("http", "https"):
        return False, "Only HTTP/HTTPS URLs are supported."
    hostname = (parsed.hostname or "").lower()
    if not hostname or "." not in hostname:
        return False, "Invalid domain."
    if hostname in _PRIVATE_HOSTNAMES or hostname.startswith(_PRIVATE_PREFIXES):
        return False, "Local/private URLs are not allowed."
    return True, url


# ===========================
# UTILS
# ===========================
def normalize_domain(url_or_domain: str) -> str:
    s = (url_or_domain or "").strip()
    if not s:
        return ""
    if s.startswith(("http://", "https://")):
        s = urlparse(s).netloc
    s = s.lower()
    if s.startswith("www."):
        s = s[4:]
    s = s.split(":")[0]
    return s


def load_prompt(path: Path) -> str:
    if path.exists():
        return path.read_text(encoding="utf-8")
    return ""


def strip_json_fences(text: str) -> str:
    t = (text or "").strip()
    t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*```$", "", t)
    return t.strip()


def safe_int(x, default=0):
    try:
        return int(x)
    except Exception:
        return default


def clean_text(text: str) -> str:
    """Decode HTML entities (including double-encoded) and strip ALL HTML tags."""
    if not text or not isinstance(text, str):
        return text or ""
    cleaned = text
    # Loop unescape to handle double/triple encoding: &amp;lt;p&amp;gt; → &lt;p&gt; → <p>
    for _ in range(3):
        prev = cleaned
        cleaned = html_module.unescape(cleaned)
        if cleaned == prev:
            break
    # Use BeautifulSoup to reliably strip ALL HTML tags (including nested/complex)
    cleaned = BeautifulSoup(cleaned, "html.parser").get_text()
    # Collapse extra whitespace
    cleaned = " ".join(cleaned.split())
    return cleaned.strip()


def safe_html(text: str) -> str:
    """Clean text AND escape it for safe injection into HTML templates.

    Use this for any text that will be placed inside an HTML template
    rendered via st.markdown(unsafe_allow_html=True). It ensures that
    characters like < > & in the text don't break the surrounding HTML.
    """
    cleaned = clean_text(text)
    return html_module.escape(cleaned)


def clean_audit_data(data):
    """Recursively clean all string values in the audit data structure."""
    if isinstance(data, str):
        return clean_text(data)
    if isinstance(data, list):
        return [clean_audit_data(item) for item in data]
    if isinstance(data, dict):
        return {k: clean_audit_data(v) for k, v in data.items()}
    return data


# ===========================
# WEB ANALYSIS (single page snapshot)
# ===========================
def analyze_basic_site(url: str) -> dict:
    try:
        response = http_session.get(url, timeout=CRAWL_TIMEOUT)
        soup = BeautifulSoup(response.content, "html.parser")
        base_domain = normalize_domain(url)

        analysis = {
            "url": url,
            "status_code": response.status_code,
            "title": soup.title.string.strip() if soup.title and soup.title.string else "No title found",
            "meta_description": "",
            "h1_tags": [],
            "h2_tags": [],
            "internal_links": 0,
            "external_links": 0,
            "word_count": 0,
        }

        meta_desc = soup.find("meta", attrs={"name": "description"})
        if meta_desc:
            analysis["meta_description"] = meta_desc.get("content", "")

        analysis["h1_tags"] = [h1.get_text(" ", strip=True) for h1 in soup.find_all("h1")]
        analysis["h2_tags"] = [h2.get_text(" ", strip=True) for h2 in soup.find_all("h2")][:5]

        for link in soup.find_all("a", href=True):
            href = link["href"].strip()
            if href.startswith(("mailto:", "tel:", "#", "javascript:")):
                continue
            if href.startswith("http"):
                if normalize_domain(href) != base_domain:
                    analysis["external_links"] += 1
                else:
                    analysis["internal_links"] += 1
            else:
                analysis["internal_links"] += 1

        text = soup.get_text(" ", strip=True)
        analysis["word_count"] = len(text.split())

        logger.info("Homepage snapshot OK: %s (status %d)", url, response.status_code)
        return analysis

    except Exception as e:
        logger.error("Homepage snapshot failed for %s: %s", url, e)
        return {"error": str(e)}


# ===========================
# LIGHTWEIGHT CRAWLER
# ===========================
def fetch_url(url: str, timeout: int = CRAWL_TIMEOUT):
    try:
        return http_session.get(url, timeout=timeout, allow_redirects=True)
    except Exception as e:
        logger.debug("fetch_url failed for %s: %s", url, e)
        return None


def _parse_robots_for_sitemaps(robots_text: str) -> list[str]:
    """Extract Sitemap: directives from robots.txt content."""
    sitemaps = []
    for line in robots_text.splitlines():
        if line.lower().startswith("sitemap:"):
            sm = line.split(":", 1)[1].strip()
            if sm:
                sitemaps.append(sm)
    return list(dict.fromkeys(sitemaps))


def get_robots_sitemaps(base_url: str) -> list[str]:
    """Check robots.txt for Sitemap directives, trying both www and non-www."""
    parsed = urlparse(base_url)
    hostname = parsed.hostname or ""

    # Build list of robots.txt URLs to try
    urls_to_try = [urljoin(base_url.rstrip("/") + "/", "robots.txt")]

    # Also try the www/non-www counterpart
    if hostname.startswith("www."):
        alt_base = f"{parsed.scheme}://{hostname[4:]}"
        urls_to_try.append(urljoin(alt_base.rstrip("/") + "/", "robots.txt"))
    elif hostname and "." in hostname:
        alt_base = f"{parsed.scheme}://www.{hostname}"
        urls_to_try.append(urljoin(alt_base.rstrip("/") + "/", "robots.txt"))

    for robots_url in urls_to_try:
        r = fetch_url(robots_url)
        if r and r.status_code < 400:
            sitemaps = _parse_robots_for_sitemaps(r.text)
            if sitemaps:
                logger.info("Found %d sitemap(s) in %s", len(sitemaps), robots_url)
                return sitemaps
    return []


def try_default_sitemaps(base_url: str) -> list[str]:
    """Generate sitemap candidate URLs for both www and non-www variants."""
    parsed = urlparse(base_url)
    hostname = parsed.hostname or ""

    # Build base URLs to try: the original, plus the www/non-www counterpart
    bases = [base_url.rstrip("/") + "/"]
    if hostname.startswith("www."):
        alt = f"{parsed.scheme}://{hostname[4:]}"
        bases.append(alt.rstrip("/") + "/")
    elif hostname and "." in hostname:
        alt = f"{parsed.scheme}://www.{hostname}"
        bases.append(alt.rstrip("/") + "/")

    # Sitemap filenames ordered by likelihood
    filenames = [
        "sitemap_index.xml",
        "sitemap.xml",
        "wp-sitemap.xml",
        "sitemap-index.xml",
        "sitemap1.xml",
    ]

    candidates = []
    seen = set()
    for base in bases:
        for fname in filenames:
            url = urljoin(base, fname)
            if url not in seen:
                seen.add(url)
                candidates.append(url)
    return candidates


def parse_sitemap_xml(xml_text: str) -> tuple[list[str], list[str]]:
    """Parse sitemap XML (with or without namespace) and return (page_urls, child_sitemap_urls)."""
    urls, sitemaps = [], []

    # Quick guard: skip obvious non-XML responses (HTML pages returned as 200)
    stripped = (xml_text or "").strip()
    if not stripped or stripped[:5].lower().startswith("<!doc") or stripped[:5].lower().startswith("<html"):
        return urls, sitemaps

    try:
        root = ET.fromstring(xml_text)
    except Exception:
        return urls, sitemaps

    def _local_name(tag: str) -> str:
        """Strip namespace prefix from tag, e.g. {http://...}loc -> loc."""
        if "}" in tag:
            return tag.split("}", 1)[1].lower()
        return tag.lower()

    root_name = _local_name(root.tag)

    if root_name == "sitemapindex":
        # Sitemap index file — collect child sitemap URLs
        for el in root.iter():
            if _local_name(el.tag) == "loc" and el.text:
                sitemaps.append(el.text.strip())
    else:
        # Regular sitemap — collect page URLs
        for el in root.iter():
            if _local_name(el.tag) == "loc" and el.text:
                urls.append(el.text.strip())

    return urls, sitemaps


def fetch_sitemap_urls(sitemap_url: str, max_urls: int = MAX_SITEMAP_URLS) -> list[str]:
    logger.info("Trying sitemap: %s", sitemap_url)
    r = fetch_url(sitemap_url)
    if not r:
        logger.info("  -> request failed (no response)")
        return []
    if r.status_code >= 400:
        logger.info("  -> HTTP %d", r.status_code)
        return []

    urls, sitemaps = parse_sitemap_xml(r.text)
    logger.info("  -> parsed: %d page URLs, %d child sitemaps", len(urls), len(sitemaps))
    all_urls = list(urls)

    for sm in sitemaps[:20]:
        time.sleep(0.15)
        all_urls.extend(fetch_sitemap_urls(sm, max_urls=max_urls))
        if len(all_urls) >= max_urls:
            break

    deduped = list(dict.fromkeys(all_urls))
    return deduped[:max_urls]


def pick_sample_urls(urls: list[str], homepage_url: str, max_pages: int = MAX_PAGES) -> list[str]:
    urls = [u for u in urls if isinstance(u, str) and u.startswith(("http://", "https://"))]
    urls = list(dict.fromkeys(urls))

    sample = []
    if homepage_url:
        sample.append(homepage_url)

    by_bucket = defaultdict(list)
    for u in urls:
        try:
            p = urlparse(u)
            path = (p.path or "/").strip("/")
            bucket = path.split("/")[0] if path else "_root"
            by_bucket[bucket].append(u)
        except Exception:
            continue

    for _bucket, lst in by_bucket.items():
        if len(sample) >= max_pages:
            break
        chosen = lst[0]
        if chosen not in sample:
            sample.append(chosen)

    for u in urls:
        if len(sample) >= max_pages:
            break
        if u not in sample:
            sample.append(u)

    return sample[:max_pages]


def extract_page_signals(url: str, base_domain: str) -> dict:
    r = fetch_url(url)
    if not r:
        return {"url": url, "final_url": url, "status": None, "error": "request_failed"}

    final_url = r.url
    status = r.status_code
    content_type = (r.headers.get("Content-Type") or "").lower()
    if "text/html" not in content_type:
        return {"url": url, "final_url": final_url, "status": status,
                "content_type": content_type, "error": "non_html"}

    soup = BeautifulSoup(r.text, "html.parser")

    title = soup.title.string.strip() if soup.title and soup.title.string else ""

    meta_desc = ""
    md = soup.find("meta", attrs={"name": "description"})
    if md:
        meta_desc = (md.get("content") or "").strip()

    canonical = ""
    canon = soup.find("link", attrs={"rel": lambda x: x and "canonical" in x.lower()})
    if canon:
        canonical = (canon.get("href") or "").strip()

    robots_meta = ""
    rm = soup.find("meta", attrs={"name": lambda x: x and x.lower() == "robots"})
    if rm:
        robots_meta = (rm.get("content") or "").strip().lower()

    h1_count = len(soup.find_all("h1"))

    text = soup.get_text(" ", strip=True)
    word_count = len(text.split())

    jsonld_tags = soup.find_all("script", attrs={"type": "application/ld+json"})
    jsonld_count = len(jsonld_tags)

    internal_links = []
    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        abs_url = href if href.startswith(("http://", "https://")) else urljoin(final_url, href)
        if normalize_domain(abs_url) == base_domain:
            internal_links.append(abs_url)
        if len(internal_links) >= MAX_INTERNAL_LINKS_PER_PAGE:
            break

    return {
        "url": url,
        "final_url": final_url,
        "status": status,
        "title": title,
        "title_len": len(title),
        "meta": meta_desc,
        "meta_len": len(meta_desc),
        "canonical": canonical,
        "robots_meta": robots_meta,
        "h1_count": h1_count,
        "word_count": word_count,
        "jsonld_count": jsonld_count,
        "sample_internal_links": internal_links,
    }


def check_links_for_broken(links: list[str]) -> tuple[int, list[dict]]:
    broken = []
    ok = 0
    for link in links:
        try:
            r = http_session.head(link, timeout=CRAWL_TIMEOUT, allow_redirects=True)
            code = r.status_code
            if code >= 400 or code == 0:
                rg = http_session.get(link, timeout=CRAWL_TIMEOUT, allow_redirects=True)
                code = rg.status_code
            if code >= 400:
                broken.append({"url": link, "status": code})
            else:
                ok += 1
        except Exception:
            broken.append({"url": link, "status": None})
        time.sleep(0.05)
        if len(broken) >= MAX_BROKEN_LINKS_REPORTED:
            break
    return ok, broken


def build_site_level_findings(pages: list[dict], base_domain: str) -> tuple[dict, dict]:
    summary = {
        "analyzed_pages": len(pages),
        "status_4xx_5xx": 0,
        "redirects": 0,
        "missing_title": 0,
        "missing_meta": 0,
        "missing_h1": 0,
        "multiple_h1": 0,
        "noindex_pages": 0,
        "missing_canonical": 0,
        "canonical_mismatch": 0,
        "thin_pages": 0,
        "pages_with_schema": 0,
    }

    examples = {
        "duplicate_titles": [],
        "duplicate_meta": [],
        "noindex_examples": [],
        "canonical_examples": [],
        "thin_examples": [],
        "status_examples": [],
    }

    titles = defaultdict(list)
    metas = defaultdict(list)

    for p in pages:
        status = p.get("status")
        url = p.get("final_url") or p.get("url")

        if status is None or (isinstance(status, int) and status >= 400):
            summary["status_4xx_5xx"] += 1
            if len(examples["status_examples"]) < 10:
                examples["status_examples"].append({"url": url, "status": status})
        elif (p.get("final_url") or p.get("url")) != p.get("url"):
            summary["redirects"] += 1

        title = (p.get("title") or "").strip()
        meta = (p.get("meta") or "").strip()

        if not title:
            summary["missing_title"] += 1
        else:
            titles[title].append(url)

        if not meta:
            summary["missing_meta"] += 1
        else:
            metas[meta].append(url)

        h1c = safe_int(p.get("h1_count"), 0)
        if h1c == 0:
            summary["missing_h1"] += 1
        elif h1c > 1:
            summary["multiple_h1"] += 1

        robots = (p.get("robots_meta") or "").lower()
        if "noindex" in robots:
            summary["noindex_pages"] += 1
            if len(examples["noindex_examples"]) < 10:
                examples["noindex_examples"].append({"url": url, "robots": robots})

        canonical = (p.get("canonical") or "").strip()
        if not canonical:
            summary["missing_canonical"] += 1
        else:
            try:
                c_abs = urljoin(url, canonical) if canonical.startswith("/") else canonical
                if normalize_domain(c_abs) != base_domain:
                    summary["canonical_mismatch"] += 1
                    if len(examples["canonical_examples"]) < 10:
                        examples["canonical_examples"].append({"url": url, "canonical": canonical})
            except Exception:
                pass

        wc = safe_int(p.get("word_count"), 0)
        if 0 < wc < THIN_CONTENT_THRESHOLD:
            summary["thin_pages"] += 1
            if len(examples["thin_examples"]) < 10:
                examples["thin_examples"].append({"url": url, "word_count": wc})

        if safe_int(p.get("jsonld_count"), 0) > 0:
            summary["pages_with_schema"] += 1

    dup_titles = sorted(
        [(t, urls) for t, urls in titles.items() if len(urls) > 1],
        key=lambda x: len(x[1]),
        reverse=True,
    )
    for t, urls in dup_titles[:5]:
        examples["duplicate_titles"].append({"value": t[:140], "count": len(urls), "urls": urls[:5]})

    dup_meta = sorted(
        [(m, urls) for m, urls in metas.items() if len(urls) > 1],
        key=lambda x: len(x[1]),
        reverse=True,
    )
    for m, urls in dup_meta[:5]:
        examples["duplicate_meta"].append({"value": m[:160], "count": len(urls), "urls": urls[:5]})

    return summary, examples


_SITEMAP_PRIORITY_BEST = ("sitemap_index.xml", "sitemap.xml", "sitemap-index.xml")
_SITEMAP_PRIORITY_OK = ("sitemap1.xml", "post-sitemap.xml", "page-sitemap.xml", "wp-sitemap.xml")
_SITEMAP_PRIORITY_AVOID = ("news-sitemap.xml", "video-sitemap.xml", "image-sitemap.xml")
MIN_USEFUL_URLS = 20


def _sitemap_sort_key(url: str) -> int:
    """Return sort key: 0=best, 1=ok, 2=avoid, 3=unknown."""
    fname = url.rsplit("/", 1)[-1].lower()
    if any(p in fname for p in _SITEMAP_PRIORITY_BEST):
        return 0
    if any(p in fname for p in _SITEMAP_PRIORITY_OK):
        return 1
    if any(p in fname for p in _SITEMAP_PRIORITY_AVOID):
        return 2
    return 3


def discover_urls_from_homepage(homepage_url: str, base_domain: str) -> list[str]:
    """Crawl the homepage and extract unique internal links as a sitemap fallback."""
    r = fetch_url(homepage_url)
    if not r or r.status_code >= 400:
        return []
    soup = BeautifulSoup(r.text, "html.parser")
    urls = []
    seen = {homepage_url, homepage_url.rstrip("/"), homepage_url + "/"}
    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        abs_url = href if href.startswith(("http://", "https://")) else urljoin(homepage_url, href)
        # Strip fragment
        abs_url = abs_url.split("#")[0]
        if not abs_url or abs_url in seen:
            continue
        if normalize_domain(abs_url) == base_domain:
            seen.add(abs_url)
            urls.append(abs_url)
    logger.info("Homepage link discovery found %d internal URLs", len(urls))
    return urls


def run_basic_audit(url_input: str) -> dict:
    base_domain = normalize_domain(url_input)
    if not base_domain:
        raise ValueError("Invalid URL/domain")

    base_url = url_input if url_input.startswith(("http://", "https://")) else "https://" + url_input
    p = urlparse(base_url)
    base_url = f"{p.scheme}://{p.netloc}"

    logger.info("Starting audit for %s", base_domain)

    # Gather sitemaps from robots.txt first, then try common default paths as fallback
    robots_sitemaps = get_robots_sitemaps(base_url)
    default_sitemaps = try_default_sitemaps(base_url)
    # Combine: robots.txt sitemaps first, then defaults (deduplicated)
    seen = set()
    sitemaps = []
    for sm in robots_sitemaps + default_sitemaps:
        if sm not in seen:
            seen.add(sm)
            sitemaps.append(sm)

    # Sort candidates by priority BEFORE fetching (best first, news last)
    sitemaps.sort(key=_sitemap_sort_key)
    logger.info("Sitemap candidates (priority order): %s", sitemaps)

    discovered_urls = []
    used_sitemap = None
    fallback_urls = []
    fallback_sitemap = None
    for sm in sitemaps:
        urls = fetch_sitemap_urls(sm, max_urls=MAX_SITEMAP_URLS)
        if not urls:
            continue
        # If this sitemap has enough URLs, use it immediately
        if len(urls) >= MIN_USEFUL_URLS:
            discovered_urls = urls
            used_sitemap = sm
            logger.info("Using sitemap %s (%d URLs)", sm, len(urls))
            break
        # Otherwise keep as fallback (first small sitemap found)
        if not fallback_urls:
            fallback_urls = urls
            fallback_sitemap = sm

    # If no sitemap had enough URLs, use the best small one we found
    if not discovered_urls and fallback_urls:
        discovered_urls = fallback_urls
        used_sitemap = fallback_sitemap
        logger.info("Using fallback sitemap %s (%d URLs)", used_sitemap, len(discovered_urls))

    homepage = base_url
    if not discovered_urls:
        # Fallback: discover internal links from the homepage
        logger.info("No sitemap found — falling back to homepage link discovery")
        discovered_urls = discover_urls_from_homepage(homepage, base_domain)
        if discovered_urls:
            discovery_method = "homepage_links (no sitemap found)"
            urls_discovered_count = len(discovered_urls)
            sample_urls = pick_sample_urls(discovered_urls, homepage, max_pages=MAX_PAGES)
        else:
            sample_urls = [homepage]
            discovery_method = "homepage_only (no sitemap found)"
            urls_discovered_count = 1
    else:
        sample_urls = pick_sample_urls(discovered_urls, homepage, max_pages=MAX_PAGES)
        discovery_method = f"robots/sitemap ({used_sitemap})"
        urls_discovered_count = len(discovered_urls)

    logger.info("Discovered %d URLs, sampling %d", urls_discovered_count, len(sample_urls))

    pages = []
    for u in sample_urls:
        pages.append(extract_page_signals(u, base_domain=base_domain))
        time.sleep(0.12)

    crawl_summary, examples = build_site_level_findings(pages, base_domain=base_domain)

    all_links = []
    for pz in pages:
        for ln in pz.get("sample_internal_links") or []:
            if ln not in all_links:
                all_links.append(ln)
            if len(all_links) >= MAX_BROKEN_LINK_CHECKS:
                break
        if len(all_links) >= MAX_BROKEN_LINK_CHECKS:
            break

    ok_count, broken_examples = check_links_for_broken(all_links)
    crawl_summary["broken_internal_links_checked"] = len(all_links)
    crawl_summary["broken_internal_links_found"] = len(broken_examples)
    examples["broken_links"] = broken_examples

    logger.info("Audit finished: %d pages analyzed, %d broken links found",
                len(pages), len(broken_examples))

    return {
        "domain": base_domain,
        "audit_date": datetime.now().strftime("%B %Y"),
        "discovery_method": discovery_method,
        "urls_discovered": urls_discovered_count,
        "urls_analyzed": len(pages),
        "crawl_summary": crawl_summary,
        "pages": pages,
        "examples": examples,
    }


# ===========================
# AI
# ===========================
def run_llm(prompt_text: str) -> str:
    model = genai.GenerativeModel("gemini-3-flash-preview")
    resp = model.generate_content(prompt_text)
    return (getattr(resp, "text", "") or "").strip()


def build_prompt(context: dict, detected_problems: dict) -> str:
    """Build LLM prompt injecting both detected problems and crawl context."""
    p = load_prompt(PROMPT_BASIC).strip()
    if not p:
        p = (
            "You are a senior SEO auditor.\n"
            "Return ONLY valid JSON with executive_summary and next_checks.\n"
            "DETECTED_PROBLEMS:\n{{DETECTED_PROBLEMS}}\n"
            "CONTEXT_JSON:\n{{CONTEXT_JSON}}\n"
        )
    # Build a concise summary of detected problems for the LLM
    problems_summary = []
    for category in ("quick_wins", "critical_errors", "warnings"):
        for prob in detected_problems.get(category, []):
            problems_summary.append({
                "title": prob.title,
                "severity": prob.severity,
                "affected_urls": len(prob.urls),
                "description": prob.description,
            })
    p = p.replace("{{DETECTED_PROBLEMS}}", json.dumps(problems_summary, ensure_ascii=False, indent=2))
    p = p.replace("{{CONTEXT_JSON}}", json.dumps(context, ensure_ascii=False, indent=2))
    return p


def parse_audit_json(raw: str) -> dict:
    """Parse the JSON response from the LLM, handling code fences and HTML entities."""
    cleaned = strip_json_fences(raw)
    result = {}
    try:
        result = json.loads(cleaned)
    except json.JSONDecodeError:
        # Try to find JSON object in the response
        match = re.search(r'\{[\s\S]*\}', cleaned)
        if match:
            try:
                result = json.loads(match.group())
            except json.JSONDecodeError:
                pass
    # Clean all string values of HTML entities and residual tags
    return clean_audit_data(result)






# ===========================
# EXCEL GENERATION
# ===========================
def create_excel_report(
    quick_wins: list,
    critical_errors: list,
    warnings: list,
    next_checks: list,
    site_name: str,
) -> BytesIO:
    """Create Excel report from Problem objects. One row per affected URL in sheets 2-3."""
    wb = Workbook()

    # Styles
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill(start_color="1e3a8a", end_color="1e3a8a", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000"),
    )
    green_fill = PatternFill(start_color="d1fae5", end_color="d1fae5", fill_type="solid")

    def style_header_row(ws, row_num, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
        ws.row_dimensions[row_num].height = 30

    def style_data_cell(ws, row_num, col_num):
        cell = ws.cell(row=row_num, column=col_num)
        cell.border = thin_border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        return cell

    def build_issues_sheet(ws, problems: list):
        """Build Critical Errors or Warnings sheet — one row per affected URL."""
        headers = ["\u2713 Done", "Issue", "URL", "Description", "Why It Matters", "How to Fix"]
        ws.append(headers)
        style_header_row(ws, 1, len(headers))

        row_num = 1
        for problem in problems:
            for url in problem.urls:
                row_num += 1
                row = [
                    "\u2610",
                    clean_text(problem.title),
                    url,
                    clean_text(problem.description),
                    clean_text(problem.why_it_matters),
                    clean_text(problem.how_to_fix),
                ]
                ws.append(row)
                for c in range(1, len(row) + 1):
                    style_data_cell(ws, row_num, c)
                ws.row_dimensions[row_num].height = 50

        # Checkbox dropdown validation on column A (data rows)
        data_rows = row_num - 1
        if data_rows > 0:
            dv = DataValidation(
                type="list",
                formula1='"\u2610,\u2611"',
                allow_blank=False,
            )
            dv.error = "Please select a valid option"
            dv.errorTitle = "Invalid input"
            dv.add(f"A2:A{row_num}")
            ws.add_data_validation(dv)

            # Conditional formatting: green row when checked
            for col_idx in range(1, len(headers) + 1):
                col_letter = get_column_letter(col_idx)
                cell_range = f"{col_letter}2:{col_letter}{row_num}"
                ws.conditional_formatting.add(
                    cell_range,
                    FormulaRule(
                        formula=[f'$A2="\u2611"'],
                        fill=green_fill,
                    ),
                )

        # Column widths
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 55
        ws.column_dimensions["D"].width = 40
        ws.column_dimensions["E"].width = 35
        ws.column_dimensions["F"].width = 40

    # --- Sheet 1: Quick Wins ---
    ws_qw = wb.active
    ws_qw.title = "Quick Wins"
    headers_qw = ["#", "Issue", "Pages Affected", "Description", "Why It Matters", "How to Fix"]
    ws_qw.append(headers_qw)
    style_header_row(ws_qw, 1, len(headers_qw))

    for i, qw in enumerate(quick_wins[:5], 1):
        row_num = i + 1
        row = [
            i,
            clean_text(qw.title),
            len(qw.urls),
            clean_text(qw.description),
            clean_text(qw.why_it_matters),
            clean_text(qw.how_to_fix),
        ]
        ws_qw.append(row)
        for c in range(1, len(row) + 1):
            style_data_cell(ws_qw, row_num, c)
        ws_qw.row_dimensions[row_num].height = 50

    ws_qw.column_dimensions["A"].width = 6
    ws_qw.column_dimensions["B"].width = 30
    ws_qw.column_dimensions["C"].width = 14
    ws_qw.column_dimensions["D"].width = 45
    ws_qw.column_dimensions["E"].width = 35
    ws_qw.column_dimensions["F"].width = 40

    # --- Sheet 2: Critical Errors ---
    ws_ce = wb.create_sheet("Critical Errors")
    build_issues_sheet(ws_ce, critical_errors)

    # --- Sheet 3: Warnings ---
    ws_w = wb.create_sheet("Warnings")
    build_issues_sheet(ws_w, warnings)

    # --- Sheet 4: Next Checks ---
    ws_nc = wb.create_sheet("Next Checks")
    headers_nc = ["Check", "Description"]
    ws_nc.append(headers_nc)
    style_header_row(ws_nc, 1, len(headers_nc))

    for i, nc in enumerate(next_checks, 1):
        row_num = i + 1
        nc_title = nc.get("title", "") if isinstance(nc, dict) else ""
        nc_desc = nc.get("description", "") if isinstance(nc, dict) else ""
        row = [clean_text(nc_title), clean_text(nc_desc)]
        ws_nc.append(row)
        for c in range(1, len(row) + 1):
            style_data_cell(ws_nc, row_num, c)
        ws_nc.row_dimensions[row_num].height = 60

    ws_nc.column_dimensions["A"].width = 35
    ws_nc.column_dimensions["B"].width = 70

    # Save
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ===========================
# DOCX GENERATION (Markdown to docx) - kept for compatibility
# ===========================
def create_word_from_content(audit_content: str, site_name: str) -> BytesIO:
    doc = Document()
    title = doc.add_heading(f"SEO Audit - {site_name}", 0)
    title.alignment = 1

    subtitle = doc.add_paragraph("Basic Audit")
    subtitle.alignment = 1
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(96, 165, 250)

    doc.add_paragraph()

    for line in (audit_content or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line), style="List Number")
        elif line == "---":
            doc.add_paragraph("_" * 60)
        else:
            doc.add_paragraph(line.replace("**", ""))

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated by Quick Wins - {datetime.now().strftime('%B %d, %Y')}")
    footer.alignment = 1
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ===========================
# UI RENDERING FUNCTIONS
# ===========================
def render_problem_expander(problem, max_urls: int = 3):
    """Render a single Problem using native Streamlit components."""
    label = f"{problem.title} — {len(problem.urls)} pages affected"
    with st.expander(label):
        st.write(problem.description)

        st.write("**Why it matters:**")
        st.write(problem.why_it_matters)

        st.write("**How to fix:**")
        st.write(problem.how_to_fix)

        st.write("**Affected URLs:**")
        for url in problem.urls[:max_urls]:
            st.code(url, language=None)
        remaining = len(problem.urls) - max_urls
        if remaining > 0:
            st.caption(f"+ {remaining} more URLs — see the downloadable Excel for the full list")


# ===========================
# SIDEBAR
# ===========================
with st.sidebar:
    st.markdown("### System Status")

    if GEMINI_AVAILABLE:
        st.markdown(
            '<span class="status-badge status-connected">&#10003; AI Connected</span>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<span class="status-badge status-disconnected">&#10007; AI Offline</span>',
            unsafe_allow_html=True,
        )

    st.markdown("---")

    st.markdown("### About")
    st.markdown("""
    **Quick Wins** generates SEO audits in seconds.

    **Features**:
    - Full site crawl & analysis
    - AI-powered insights
    - Actionable recommendations
    - Downloadable Excel reports
    """)

    st.markdown("---")
    st.caption("v5.0")


# ===========================
# MAIN INTERFACE
# ===========================
st.markdown("""
<div class="app-header">
    <div class="app-title">QUICK WINS</div>
    <div class="app-subtitle">Instant AI-powered SEO audits. Crawl any site, detect issues, and get actionable recommendations in seconds.</div>
</div>
""", unsafe_allow_html=True)

# Info box
st.info("**SEO Audit** — Crawl your site, analyze pages, detect critical issues, and get AI-powered recommendations.")

st.markdown("---")

# URL Input
url_input = st.text_input(
    "Website URL",
    placeholder="example.com",
    help="Enter any URL or domain — e.g. example.com, www.example.com, https://example.com"
)

st.markdown("---")

# Generate Button
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    if st.button("Audit Now!", disabled=not url_input, use_container_width=True):
        # Validate URL
        is_valid, result = validate_url(url_input)
        if not is_valid:
            st.error(f"Invalid URL: {result}")
            st.stop()
        clean_url = result

        st.markdown("---")

        progress_bar = st.progress(0)
        status_text = st.empty()

        domain = normalize_domain(clean_url)
        site_name = domain or clean_url.replace("https://", "").replace("http://", "").replace("www.", "").split("/")[0]

        # ── Step 1: Crawl ──
        status_text.text("Discovering pages (robots/sitemap) and sampling URLs...")
        progress_bar.progress(15)

        try:
            audit_context = run_basic_audit(clean_url)
        except Exception as e:
            logger.error("Audit crawl failed: %s", e)
            st.error(f"Audit crawl failed: {e}")
            st.stop()

        progress_bar.progress(40)

        # ── Step 2: Deterministic detection ──
        status_text.text("Analyzing pages and detecting issues...")
        progress_bar.progress(50)

        pages = audit_context.get("pages", [])
        broken_examples = audit_context.get("examples", {}).get("broken_links", [])

        detected = detect_problems(
            pages=pages,
            broken_link_examples=broken_examples,
            thin_content_threshold=THIN_CONTENT_THRESHOLD,
        )

        quick_wins = detected["quick_wins"]
        critical_errors = detected["critical_errors"]
        warnings_list = detected["warnings"]

        # ── Step 3: Validate results ──
        status_text.text("Validating results...")
        progress_bar.progress(60)

        try:
            validate_results(quick_wins, critical_errors, warnings_list)
        except ValueError as e:
            logger.error("Validation failed: %s", e)
            st.error(f"Results validation failed. Please report this issue.\n\n{e}")
            st.stop()

        # ── Step 4: AI enrichment (executive summary + next checks) ──
        exec_summary = ""
        next_checks = []

        if GEMINI_AVAILABLE:
            status_text.text("Generating AI-powered summary...")
            progress_bar.progress(70)

            try:
                prompt_text = build_prompt(audit_context, detected)
                raw_response = run_llm(prompt_text)
                llm_data = parse_audit_json(raw_response)
                exec_summary = clean_text(llm_data.get("executive_summary", ""))
                next_checks = llm_data.get("next_checks", [])
            except Exception as e:
                logger.warning("AI enrichment failed (non-fatal): %s", e)
                exec_summary = (
                    f"Audited {audit_context.get('urls_analyzed', 0)} pages from {site_name} "
                    f"(discovered {audit_context.get('urls_discovered', 0)} via "
                    f"{audit_context.get('discovery_method', 'sitemap')}). "
                    f"Found {len(critical_errors)} critical issues and {len(warnings_list)} warnings."
                )
        else:
            exec_summary = (
                f"Audited {audit_context.get('urls_analyzed', 0)} pages from {site_name} "
                f"(discovered {audit_context.get('urls_discovered', 0)} via "
                f"{audit_context.get('discovery_method', 'sitemap')}). "
                f"Found {len(critical_errors)} critical issues and {len(warnings_list)} warnings."
            )

        # ── Step 5: Create Excel ──
        status_text.text("Creating report...")
        progress_bar.progress(85)

        excel_file = create_excel_report(
            quick_wins=quick_wins,
            critical_errors=critical_errors,
            warnings=warnings_list,
            next_checks=next_checks,
            site_name=site_name,
        )

        progress_bar.progress(100)
        status_text.text("Complete!")
        time.sleep(0.5)

        progress_bar.empty()
        status_text.empty()

        # ===========================
        # RESULTS DISPLAY (native Streamlit components)
        # ===========================

        # --- Audit Banner ---
        st.markdown(f"""
        <div class="audit-banner">
            <h2>SEO AUDIT — {safe_html(site_name)}</h2>
        </div>
        """, unsafe_allow_html=True)

        # --- Executive Summary ---
        if exec_summary:
            st.markdown(f"""
            <div class="info-box">
                <h3>Executive Summary</h3>
                <p>{safe_html(exec_summary)}</p>
            </div>
            """, unsafe_allow_html=True)

        # --- Audit Scope (from crawl context, not LLM) ---
        scope_html = f"""
        <div class="info-box">
            <h3>Audit Scope &amp; Method</h3>
            <p><strong>URLs discovered:</strong> {audit_context.get('urls_discovered', 'N/A')}</p>
            <p><strong>URLs analyzed:</strong> {audit_context.get('urls_analyzed', 'N/A')}</p>
            <p><strong>Discovery method:</strong> {safe_html(str(audit_context.get('discovery_method', 'N/A')))}</p>
        </div>
        """
        st.markdown(scope_html, unsafe_allow_html=True)

        # --- Quick Wins (top 3 in styled container) ---
        top_qw = quick_wins[:3]
        top_qw_titles = {qw.title for qw in top_qw}

        if top_qw:
            qw_cards_html = ""
            for i, qw in enumerate(top_qw, 1):
                urls_html = ""
                for url in qw.urls[:3]:
                    urls_html += f'<div>{url}</div>'
                remaining = len(qw.urls) - 3
                more_html = ""
                if remaining > 0:
                    more_html = f'<div class="qw-more">+{remaining} more — see Excel for full list</div>'

                qw_cards_html += f"""
                <div class="qw-card">
                    <span class="qw-number">{i}</span>
                    <span class="qw-title">{safe_html(qw.title)} — {len(qw.urls)} pages</span>
                    <div class="qw-desc">{safe_html(qw.description)}</div>
                    <div class="qw-urls">{urls_html}</div>
                    {more_html}
                </div>
                """

            st.markdown(f"""
            <div class="qw-container">
                <div class="qw-container-title">&#9889; Quick Wins</div>
                <div class="qw-container-subtitle">The most impactful improvements you can make right now</div>
                {qw_cards_html}
            </div>
            """, unsafe_allow_html=True)

        # --- Remaining errors (excluding ones already shown as Quick Wins) ---
        remaining_critical = [p for p in critical_errors if p.title not in top_qw_titles]
        remaining_warnings = [p for p in warnings_list if p.title not in top_qw_titles]

        if remaining_critical:
            total_urls = sum(len(p.urls) for p in remaining_critical)
            st.header("Critical Errors")
            st.caption(f"{len(remaining_critical)} additional critical issues affecting {total_urls} URLs")

            for err in remaining_critical:
                render_problem_expander(err)

        if remaining_warnings:
            total_urls = sum(len(p.urls) for p in remaining_warnings)
            st.header("Warnings")
            st.caption(f"{len(remaining_warnings)} warnings affecting {total_urls} URLs")

            for warn in remaining_warnings:
                render_problem_expander(warn)

        # --- Next Checks ---
        if next_checks:
            st.markdown('<div class="section-header">&#128270; NEXT CHECKS</div>', unsafe_allow_html=True)
            st.markdown('<p class="section-subtitle">Deeper analysis to unlock more improvements</p>', unsafe_allow_html=True)

            for nc in next_checks:
                nc_title = safe_html(nc.get('title', '')) if isinstance(nc, dict) else ''
                nc_desc = safe_html(nc.get('description', '')) if isinstance(nc, dict) else ''
                st.markdown(f"""
                <div class="next-card">
                    <h4>{nc_title}</h4>
                    <p>{nc_desc}</p>
                </div>
                """, unsafe_allow_html=True)

        # --- No issues found ---
        if not quick_wins and not critical_errors and not warnings_list:
            st.success(
                "No significant SEO issues detected in the sampled pages. "
                "This is a good sign! Consider running a deeper crawl for a more comprehensive audit."
            )

        # --- Action Buttons ---
        st.markdown("---")

        btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 1])

        with btn_col1:
            st.download_button(
                label="Download Full Report (Excel)",
                data=excel_file,
                file_name=f"SEO_Audit_{site_name}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        with btn_col2:
            if st.button("New Audit", use_container_width=True):
                st.rerun()
